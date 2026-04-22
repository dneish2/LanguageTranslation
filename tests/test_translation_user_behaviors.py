import os
import sys
import time
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from TranslationBackend import TranslationBackend


class DummyProgress:
    def __init__(self) -> None:
        self.value = 0

    def set_value(self, value):
        self.value = value


class DummyLabel:
    def __init__(self) -> None:
        self.text = ""


def _docx_bytes(*paragraphs: str) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def test_cancel_halts_processing_progression(monkeypatch):
    """User cancel should stop further segment processing."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    backend = TranslationBackend()
    progress = DummyProgress()
    label = DummyLabel()

    calls = {"count": 0}

    def translate_and_cancel(text, target_language, **_kwargs):
        calls["count"] += 1
        backend.request_cancel()
        return f"translated:{text}"

    monkeypatch.setattr(backend, "translate_text", translate_and_cancel)
    monkeypatch.setattr(backend, "calculate_tokens", lambda _text: 0)

    docx_stream = BytesIO(_docx_bytes("First", "Second", "Third"))

    _out, processed, _tokens, _text, segment_map = backend.process_docx(
        docx_stream,
        target_language="Spanish",
        progress_ui=progress,
        label_ui=label,
        do_translate=True,
    )

    assert calls["count"] == 1
    assert processed == 1
    assert len(segment_map) == 1
    assert progress.value < 100


def test_unsupported_extension_fails_clearly(monkeypatch):
    """User gets an explicit error when file format is unsupported."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    backend = TranslationBackend()

    with pytest.raises(ValueError, match=r"Unsupported file extension: xlsx"):
        backend.translate_file(
            input_stream=BytesIO(b"unused"),
            file_extension="xlsx",
            target_language="French",
            progress_ui=DummyProgress(),
            label_ui=DummyLabel(),
        )


def test_segment_update_modifies_regenerated_output(monkeypatch):
    """User edits should appear in the regenerated downloadable file."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    backend = TranslationBackend()
    monkeypatch.setattr(backend, "calculate_tokens", lambda _text: 0)

    source_stream = BytesIO(_docx_bytes("Original paragraph"))
    backend.process_docx(
        source_stream,
        target_language="Spanish",
        progress_ui=DummyProgress(),
        label_ui=DummyLabel(),
        do_translate=False,
    )

    seg_id, segment = next(iter(backend.segment_map.items()))
    assert segment["type"] == "paragraph"

    backend.update_segment(seg_id, "Contenido actualizado", target_language="Spanish")

    regenerated = Document(BytesIO(backend.output_stream.getvalue()))
    assert regenerated.paragraphs[0].text == "Contenido actualizado"


def test_translate_fallback_returns_original_on_openai_exception(monkeypatch):
    """If OpenAI fails, user still receives their original text."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    backend = TranslationBackend()

    def raise_openai_error(**_kwargs):
        raise RuntimeError("simulated OpenAI outage")

    monkeypatch.setattr(backend.client.chat.completions, "create", raise_openai_error)

    original = "\tKeep me as-is\t"
    translated = backend.translate_text(original, target_language="German")

    assert translated == "Keep me as-is"


def test_translation_job_reaches_succeeded_with_result_handle(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    backend = TranslationBackend()

    def fake_translate_file(**kwargs):
        progress_callback = kwargs.get("progress_callback")
        if progress_callback:
            progress_callback(25, "Queued for processing")
            progress_callback(90, "Almost done")
        return BytesIO(b"done"), 3, 1200, "translated text", {"seg-1": {"original": "A", "translated": "B"}}

    monkeypatch.setattr(backend, "translate_file", fake_translate_file)

    job_id = backend.start_translation_job(
        input_stream=BytesIO(b"input"),
        file_extension="docx",
        target_language="Spanish",
        processed=False,
    )

    deadline = time.time() + 2
    job = backend.get_job(job_id)
    while job and job.state in {"queued", "running"} and time.time() < deadline:
        time.sleep(0.01)
        job = backend.get_job(job_id)

    assert job is not None
    assert job.state == "succeeded"
    assert job.result_handle is not None
    result = backend.get_job_result(job.result_handle)
    assert result is not None
    assert result["count"] == 3
    assert result["segment_map"]["seg-1"]["translated"] == "B"


def test_translation_job_cancel_transitions_to_canceled(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    backend = TranslationBackend()

    def slow_translate_file(**kwargs):
        job_id = kwargs.get("job_id")
        progress_callback = kwargs.get("progress_callback")
        for idx in range(1, 20):
            if backend._is_cancel_requested(job_id):
                break
            if progress_callback:
                progress_callback(idx * 5, f"Step {idx}")
            time.sleep(0.01)
        return BytesIO(b"done"), 1, 0, "", {}

    monkeypatch.setattr(backend, "translate_file", slow_translate_file)

    job_id = backend.start_translation_job(
        input_stream=BytesIO(b"input"),
        file_extension="docx",
        target_language="Spanish",
    )
    time.sleep(0.03)
    canceled = backend.cancel_job(job_id)
    assert canceled is True

    deadline = time.time() + 2
    job = backend.get_job(job_id)
    while job and job.state in {"queued", "running"} and time.time() < deadline:
        time.sleep(0.01)
        job = backend.get_job(job_id)

    assert job is not None
    assert job.state == "canceled"


def test_cancel_one_job_does_not_cancel_other_concurrent_job(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    backend = TranslationBackend()

    def slow_translate_file(**kwargs):
        job_id = kwargs.get("job_id")
        progress_callback = kwargs.get("progress_callback")
        for idx in range(1, 20):
            if backend._is_cancel_requested(job_id):
                return BytesIO(b"canceled"), 0, 0, "", {}
            if progress_callback:
                progress_callback(idx * 5, f"Step {idx}")
            time.sleep(0.01)
        return BytesIO(f"done:{job_id}".encode()), 1, 0, f"text:{job_id}", {"seg-1": {"translated": job_id}}

    monkeypatch.setattr(backend, "translate_file", slow_translate_file)

    canceled_job_id = backend.start_translation_job(
        input_stream=BytesIO(b"input-a"),
        file_extension="docx",
        target_language="Spanish",
    )
    surviving_job_id = backend.start_translation_job(
        input_stream=BytesIO(b"input-b"),
        file_extension="docx",
        target_language="French",
    )

    time.sleep(0.03)
    assert backend.cancel_job(canceled_job_id) is True

    deadline = time.time() + 2
    canceled_job = backend.get_job(canceled_job_id)
    surviving_job = backend.get_job(surviving_job_id)
    while (
        canceled_job
        and surviving_job
        and (canceled_job.state in {"queued", "running"} or surviving_job.state in {"queued", "running"})
        and time.time() < deadline
    ):
        time.sleep(0.01)
        canceled_job = backend.get_job(canceled_job_id)
        surviving_job = backend.get_job(surviving_job_id)

    assert canceled_job is not None
    assert canceled_job.state == "canceled"
    assert surviving_job is not None
    assert surviving_job.state == "succeeded"
